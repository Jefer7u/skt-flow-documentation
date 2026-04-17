import streamlit as st
import json
import pandas as pd
import io
import os
import re
from collections import Counter
from datetime import datetime
from openpyxl.styles import PatternFill, Font, Border, Side, Alignment
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(page_title="FlowDocs · Simetrik", page_icon="📄", layout="wide")

# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════
if "lang" not in st.session_state:
    st.session_state.lang = "en"
if "dark" not in st.session_state:
    st.session_state.dark = False

# ══════════════════════════════════════════════════════════════════════════════
# STRINGS — all UI and document text, bilingual
# ══════════════════════════════════════════════════════════════════════════════
T = {
    "en": {
        # UI
        "header_title":            "Simetrik Flow Documentation",
        "header_sub":              "Reconciliation Flow Analyzer · v3.0",
        "upload_label":            "**Upload your Simetrik JSON export**",
        "upload_help":             "In Simetrik: Flow → ⚙️ Settings → Export JSON",
        "upload_placeholder_h":    "Drag the JSON here or use the button to select it",
        "upload_placeholder_p":    "In Simetrik: Flow → Settings → Export JSON",
        "step1_title":             "1️⃣  Select resources to document",
        "filter_placeholder":      "Filter by resource type…",
        "btn_all":                 "✅ All",
        "btn_none":                "☐ None",
        "btn_generate":            "🚀  GENERATE REPORTS",
        "success_msg":             "Reports generated with **{n}** resources documented.",
        "btn_excel":               "📊  Download Excel (technical)",
        "btn_word":                "📄  Download Word (executive)",
        "no_sel_warning":          "Select at least one resource to continue.",
        "sel_label":               "selected",
        "resources_n":             "resources",
        "footer":                  "Simetrik Flow Docs · v3.0",
        "lang_label":              "Language / Idioma",
        # Disclaimers
        "disc_privacy_title":  "🔒 Your data never leaves this session",
        "disc_privacy_body":   (
            "Uploaded files are processed <strong>entirely in memory</strong> and are never written to disk, "
            "stored in a database, or transmitted to any external server. "
            "When you close this tab, all data is discarded automatically. "
            "<a href='https://github.com' target='_blank' style='color:#EA0050'>View source · PRIVACY.md</a>"
        ),
        "disc_output_title":   "⚠️ Disclaimer — generated output",
        "disc_output_body":    (
            "The Excel and Word files are generated <strong>algorithmically</strong> from the uploaded JSON. "
            "Accuracy is not guaranteed. The authors are not liable for decisions made based on this output. "
            "Always verify against the original platform configuration. "
            "By downloading, you confirm you are authorized to process this data through a third-party tool."
        ),
        "disc_footer_1":       "No data is stored, transmitted, or logged — all processing happens in your browser session.",
        "disc_footer_2":       "Output is provided for informational purposes only. The authors accept no liability for its accuracy or use.",
        "disc_footer_3":       "Hosted on Streamlit Community Cloud. See PRIVACY.md for full policy.",
        "theme_label":             "Dark mode",
        # Metrics
        "m_total":    "Total",
        "m_sources":  "Sources",
        "m_unions":   "Unions",
        "m_groups":   "Aggregations",
        "m_std":      "Std. Recon.",
        "m_adv":      "Adv. Recon.",
        "m_json":     "JSON loaded",
        # Cards
        "origin":     "— origin",
        "end_flow":   "— end of flow",
        # RT labels
        "rt_native":                  "📥 Source",
        "rt_source_union":            "🔗 Source Union",
        "rt_source_group":            "📊 Group By",
        "rt_reconciliation":          "⚖️ Standard Reconciliation",
        "rt_advanced_reconciliation": "🔬 Advanced Reconciliation",
        "rt_consolidation":           "🗂️ Consolidation",
        "rt_resource_join":           "🔀 Resource Join",
        "rt_cumulative_balance":      "📈 Cumulative Balance",
        # Excel index
        "xl_idx_title":    "SIMETRIK FLOW DOCUMENTATION",
        "xl_idx_subtitle": "Technical Reference Report",
        "xl_col_num":   "#",
        "xl_col_id":    "ID",
        "xl_col_name":  "RESOURCE NAME",
        "xl_col_type":  "TYPE",
        "xl_col_from":  "COMES FROM",
        "xl_col_to":    "FEEDS TO",
        "xl_col_link":  "LINK 🔗",
        "xl_link_txt":  "View →",
        "xl_origin":    "— origin",
        "xl_end":       "— end of flow",
        # Excel meta rows
        "xl_meta_id":   "Resource ID",
        "xl_meta_type": "Type",
        "xl_meta_from": "Comes from",
        "xl_meta_to":   "Feeds to",
        # Excel section titles
        "xl_s_std":      "⚖️  STANDARD RECONCILIATION RULES",
        "xl_s_std_g":    "  ACTIVE RECONCILABLE GROUPS",
        "xl_s_std_rs":   "  MATCHING RULE SETS",
        "xl_s_adv":      "🔬  ADVANCED RECONCILIATION RULES",
        "xl_s_adv_g":    "  RECONCILABLE GROUPS & INTERNAL SEGMENTS",
        "xl_s_adv_rs":   "  RULE SETS (SEGMENT A vs SEGMENT B)",
        "xl_s_gb":       "📊  GROUP BY CONFIGURATION",
        "xl_s_union":    "🔗  SOURCE UNION CONFIGURATION",
        "xl_s_segs":     "🔍  RESOURCE RECONCILABLE GROUPS",
        "xl_s_cols":     "📋  COLUMN CONFIGURATION",
        # Excel column headers
        "xl_h_side":     "SIDE",
        "xl_h_resource": "RESOURCE",
        "xl_h_group":    "RECONCILABLE GROUP (ACTIVE)",
        "xl_h_filters":  "GROUP FILTERS",
        "xl_h_pos":      "POS.",
        "xl_h_rsname":   "RULE SET NAME",
        "xl_h_rules":    "RULES (A vs B)",
        "xl_h_seg_a":    "SEGMENT SIDE A",
        "xl_h_seg_b":    "SEGMENT SIDE B",
        "xl_h_intseg":   "INTERNAL SEGMENTS",
        "xl_h_gname":    "GROUP NAME",
        "xl_h_fapplied": "FILTERS APPLIED",
        "xl_h_usedin":   "USED IN",
        "xl_h_label":    "LABEL / NAME",
        "xl_h_dtype":    "DATA TYPE",
        "xl_h_ctype":    "COL. TYPE",
        "xl_h_logic":    "LOGIC · FORMULA · VLOOKUP",
        "xl_h_source":   "SOURCE",
        "xl_h_role":     "ROLE",
        "xl_h_gbdims":   "GROUP BY (dimensions)",
        "xl_h_aggs":     "Aggregations (metrics)",
        "xl_h_accum":    "Accumulative",
        "xl_chained":    "Chained reconciliation",
        "xl_yes":        "Yes",
        "xl_no":         "No",
        "xl_no_filters": "No filters configured",
        "xl_no_intseg":  "(no internal segmentation)",
        "xl_trigger":    "  [TRIGGER]",
        "xl_no_usage":   "Not used in active flow",
        "xl_add_src":    "Additional source",
        "xl_trig_src":   "TRIGGER · ",
        "xl_direct":     "Direct field / inherited",
        "xl_dup_bool":   "TYPE: Duplicate boolean",
        "xl_dup_int":    "TYPE: Duplicate counter",
        "xl_order_by":   "ORDER BY: ",
        "xl_part_by":    "PARTITION BY (duplicate key):\n  ",
        "xl_vlookup":    "VLOOKUP IN: ",
        "xl_match_key":  "MATCH KEY: ",
        "xl_formula":    "FORMULA: ",
        # Word doc
        "wd_title":         "Executive Flow Documentation",
        "wd_subtitle":      "Simetrik Reconciliation Flow — Configuration Summary",
        "wd_generated":     "Generated",
        "wd_res_count":     "Resources documented",
        "wd_s1":            "Executive Overview",
        "wd_s2":            "Flow Map",
        "wd_s2_note":       "Reading order: left to right, following the data pipeline. Each resource is listed with its direct inputs and outputs.",
        "wd_s2_col":        ["Resource", "Type", "What it does"],
        "wd_s3":            "Resource Breakdown",
        "wd_s4":            "Observations & Alerts",
        "wd_obs_none":      "No anomalies detected in this flow configuration.",
        "wd_obs_no_child":  "'{name}' has no downstream consumers within the documented selection — it may be a terminal output or an isolated resource.",
        "wd_obs_zero_tol":  "Rule set '{rs}' in '{res}' uses zero tolerance — any amount difference results in an unmatched item.",
        "wd_obs_chained":   "'{name}' uses chained reconciliation — ensure upstream matching is complete before this resource runs.",
        "wd_obs_unused_seg":"Segment '{seg}' (in '{res}') is defined but not referenced by any reconciliation or union in the documented resources.",
        # Word per-type templates (used in describe_resource_text)
        "wd_native_intro":  "{name} is a native data source that ingests raw records directly into the flow.",
        "wd_native_segs":   "It defines {n} reconcilable group(s) for downstream matching.",
        "wd_native_out":    "Output: feeds {children}.",
        "wd_native_no_out": "Output: no downstream consumers detected in the current selection.",
        "wd_group_intro":   "{name} applies a GROUP BY aggregation over records from {parents}.",
        "wd_group_dims":    "Grouping dimensions: {dims}.",
        "wd_group_aggs":    "Computed aggregations: {aggs}.",
        "wd_group_accum":   "The aggregation is accumulative across time periods.",
        "wd_group_out":     "Output feeds: {children}.",
        "wd_union_intro":   "{name} merges {n} source segment(s) into a unified dataset.",
        "wd_union_trigger": "Trigger side: '{seg}' from {res}.",
        "wd_union_extra":   "Additional sources: {names}.",
        "wd_union_out":     "Output feeds: {children}.",
        "wd_recon_intro":   "{name} performs a standard reconciliation between two sides:",
        "wd_recon_side":    "Side {prefix}{trigger}: resource '{res}', group '{seg}'.",
        "wd_recon_rs_n":    "Matching rule sets: {n}.",
        "wd_recon_rs":      "[{pos}] {name}: {rules}.",
        "wd_recon_chained": "This reconciliation is chained — it depends on a prior matching pass.",
        "wd_recon_out":     "Output feeds: {children}.",
        "wd_adv_intro":     "{name} performs an advanced reconciliation across {n} group(s):",
        "wd_adv_group":     "Side {prefix}: resource '{res}', group '{seg}', segmented by '{col}' ({vals}).",
        "wd_adv_no_int":    "no internal segmentation",
        "wd_adv_rs_n":      "Rule sets: {n}.",
        "wd_adv_rs":        "[{pos}] {name}: {a} vs {b}.",
        "wd_adv_out":       "Output feeds: {children}.",
        "wd_generic_intro": "{name} is a {rt} resource.",
        "wd_generic_io":    "Inputs: {parents}. Outputs: {children}.",
        "wd_no_parents":    "no upstream sources",
        "wd_no_children":   "no downstream consumers",
    },
    "es": {
        # UI
        "header_title":            "Simetrik Flow Documentation",
        "header_sub":              "Analizador de Flujos de Conciliación · v3.0",
        "upload_label":            "**Cargá el JSON exportado desde Simetrik**",
        "upload_help":             "En Simetrik: Flujo → ⚙️ Configuración → Exportar JSON",
        "upload_placeholder_h":    "Arrastrá el JSON aquí o usá el botón para seleccionarlo",
        "upload_placeholder_p":    "En Simetrik: Flujo → Configuración → Exportar JSON",
        "step1_title":             "1️⃣  Seleccioná los recursos a documentar",
        "filter_placeholder":      "Filtrá por tipo de recurso…",
        "btn_all":                 "✅ Todos",
        "btn_none":                "☐ Ninguno",
        "btn_generate":            "🚀  GENERAR REPORTES",
        "success_msg":             "Reportes generados con **{n}** recursos documentados.",
        "btn_excel":               "📊  Descargar Excel (técnico)",
        "btn_word":                "📄  Descargar Word (ejecutivo)",
        "no_sel_warning":          "Seleccioná al menos un recurso para continuar.",
        "sel_label":               "seleccionados",
        "resources_n":             "recursos",
        "footer":                  "Simetrik Flow Docs · v3.0",
        "lang_label":              "Language / Idioma",
        # Disclaimers
        "disc_privacy_title":  "🔒 Tus datos no salen de esta sesión",
        "disc_privacy_body":   (
            "Los archivos subidos se procesan <strong>íntegramente en memoria</strong> y nunca se escriben a disco, "
            "se almacenan en una base de datos ni se transmiten a ningún servidor externo. "
            "Al cerrar esta pestaña, todos los datos se descartan automáticamente. "
            "<a href='https://github.com' target='_blank' style='color:#EA0050'>Ver código · PRIVACY.md</a>"
        ),
        "disc_output_title":   "⚠️ Aviso — contenido generado",
        "disc_output_body":    (
            "Los archivos Excel y Word se generan <strong>algorítmicamente</strong> a partir del JSON subido. "
            "No se garantiza su exactitud. Los autores no son responsables de las decisiones tomadas en base a este output. "
            "Verificá siempre contra la configuración original de la plataforma. "
            "Al descargar, confirmás que estás autorizado a procesar estos datos a través de una herramienta de terceros."
        ),
        "disc_footer_1":       "No se almacenan, transmiten ni registran datos — todo el procesamiento ocurre en tu sesión de navegador.",
        "disc_footer_2":       "El output se provee solo con fines informativos. Los autores no aceptan responsabilidad por su exactitud o uso.",
        "disc_footer_3":       "Alojado en Streamlit Community Cloud. Ver PRIVACY.md para la política completa.",
        "theme_label":             "Modo oscuro",
        # Metrics
        "m_total":    "Total",
        "m_sources":  "Fuentes",
        "m_unions":   "Uniones",
        "m_groups":   "Agrupaciones",
        "m_std":      "Conc. Std",
        "m_adv":      "Conc. Avz",
        "m_json":     "JSON cargado",
        # Cards
        "origin":     "— origen",
        "end_flow":   "— fin de flujo",
        # RT labels
        "rt_native":                  "📥 Fuente",
        "rt_source_union":            "🔗 Unión de Fuentes",
        "rt_source_group":            "📊 Agrupación (Group By)",
        "rt_reconciliation":          "⚖️ Conciliación Estándar",
        "rt_advanced_reconciliation": "🔬 Conciliación Avanzada",
        "rt_consolidation":           "🗂️ Consolidación",
        "rt_resource_join":           "🔀 Join de Recursos",
        "rt_cumulative_balance":      "📈 Balance Acumulado",
        # Excel index
        "xl_idx_title":    "SIMETRIK FLOW DOCUMENTATION",
        "xl_idx_subtitle": "Reporte de Referencia Técnica",
        "xl_col_num":   "#",
        "xl_col_id":    "ID",
        "xl_col_name":  "NOMBRE DEL RECURSO",
        "xl_col_type":  "TIPO",
        "xl_col_from":  "PROVIENE DE",
        "xl_col_to":    "ALIMENTA A",
        "xl_col_link":  "LINK 🔗",
        "xl_link_txt":  "Ver →",
        "xl_origin":    "— origen",
        "xl_end":       "— fin de flujo",
        # Excel meta rows
        "xl_meta_id":   "ID Recurso",
        "xl_meta_type": "Tipo",
        "xl_meta_from": "Proviene de",
        "xl_meta_to":   "Alimenta a",
        # Excel section titles
        "xl_s_std":      "⚖️  REGLAS DE CONCILIACIÓN ESTÁNDAR",
        "xl_s_std_g":    "  GRUPOS CONCILIABLES ACTIVOS",
        "xl_s_std_rs":   "  RULE SETS DE MATCHING",
        "xl_s_adv":      "🔬  REGLAS DE CONCILIACIÓN AVANZADA",
        "xl_s_adv_g":    "  GRUPOS CONCILIABLES Y SEGMENTOS INTERNOS",
        "xl_s_adv_rs":   "  RULE SETS (SEGMENTO A vs SEGMENTO B)",
        "xl_s_gb":       "📊  CONFIGURACIÓN DE AGRUPACIÓN (GROUP BY)",
        "xl_s_union":    "🔗  CONFIGURACIÓN DE UNIÓN DE FUENTES",
        "xl_s_segs":     "🔍  GRUPOS CONCILIABLES DEL RECURSO",
        "xl_s_cols":     "📋  CONFIGURACIÓN DE COLUMNAS",
        # Excel column headers
        "xl_h_side":     "LADO",
        "xl_h_resource": "RECURSO",
        "xl_h_group":    "GRUPO CONCILIABLE (ACTIVO)",
        "xl_h_filters":  "FILTROS DEL GRUPO",
        "xl_h_pos":      "POS.",
        "xl_h_rsname":   "NOMBRE DEL RULE SET",
        "xl_h_rules":    "REGLAS  (A vs B)",
        "xl_h_seg_a":    "SEGMENTO LADO A",
        "xl_h_seg_b":    "SEGMENTO LADO B",
        "xl_h_intseg":   "SEGMENTOS INTERNOS",
        "xl_h_gname":    "NOMBRE DEL GRUPO",
        "xl_h_fapplied": "FILTROS APLICADOS",
        "xl_h_usedin":   "USADO EN",
        "xl_h_label":    "LABEL / NOMBRE",
        "xl_h_dtype":    "TIPO DATO",
        "xl_h_ctype":    "TIPO COL.",
        "xl_h_logic":    "LÓGICA · FÓRMULA · BUSCAR V",
        "xl_h_source":   "FUENTE",
        "xl_h_role":     "ROL",
        "xl_h_gbdims":   "GROUP BY (dimensiones)",
        "xl_h_aggs":     "Agregaciones (métricas)",
        "xl_h_accum":    "Acumulativo",
        "xl_chained":    "Conciliación encadenada",
        "xl_yes":        "Sí",
        "xl_no":         "No",
        "xl_no_filters": "Sin filtros configurados",
        "xl_no_intseg":  "(sin segmentación interna)",
        "xl_trigger":    "  [TRIGGER]",
        "xl_no_usage":   "Sin uso en flujo activo",
        "xl_add_src":    "Fuente adicional",
        "xl_trig_src":   "TRIGGER · ",
        "xl_direct":     "Campo directo / heredado",
        "xl_dup_bool":   "TIPO: Booleano de duplicado",
        "xl_dup_int":    "TIPO: Numeracion de duplicado",
        "xl_order_by":   "ORDER BY: ",
        "xl_part_by":    "PARTITION BY (clave de duplicado):\n  ",
        "xl_vlookup":    "BUSCAR V EN: ",
        "xl_match_key":  "CLAVE MATCH: ",
        "xl_formula":    "FÓRMULA: ",
        "xl_direct":     "Campo directo / heredado",
        # Word doc
        "wd_title":         "Documentación Ejecutiva del Flujo",
        "wd_subtitle":      "Flujo de Conciliación Simetrik — Resumen de Configuración",
        "wd_generated":     "Generado",
        "wd_res_count":     "Recursos documentados",
        "wd_s1":            "Resumen Ejecutivo",
        "wd_s2":            "Mapa del Flujo",
        "wd_s2_note":       "Orden de lectura: izquierda a derecha, siguiendo el pipeline de datos. Cada recurso se lista con sus entradas y salidas directas.",
        "wd_s2_col":        ["Recurso", "Tipo", "Qué hace"],
        "wd_s3":            "Detalle por Recurso",
        "wd_s4":            "Observaciones y Alertas",
        "wd_obs_none":      "No se detectaron anomalías en la configuración de este flujo.",
        "wd_obs_no_child":  "'{name}' no tiene consumidores downstream dentro de la selección documentada — puede ser una salida terminal o un recurso aislado.",
        "wd_obs_zero_tol":  "El rule set '{rs}' en '{res}' usa tolerancia cero — cualquier diferencia de monto resulta en un ítem sin conciliar.",
        "wd_obs_chained":   "'{name}' usa conciliación encadenada — asegurate de que el matching upstream esté completo antes de que este recurso se ejecute.",
        "wd_obs_unused_seg":"El segmento '{seg}' (en '{res}') está definido pero no es referenciado por ninguna conciliación o unión en los recursos documentados.",
        # Word per-type templates
        "wd_native_intro":  "{name} es una fuente de datos nativa que ingesta registros directamente en el flujo.",
        "wd_native_segs":   "Define {n} grupo(s) conciliable(s) para matching downstream.",
        "wd_native_out":    "Salida: alimenta a {children}.",
        "wd_native_no_out": "Salida: no se detectan consumidores downstream en la selección actual.",
        "wd_group_intro":   "{name} aplica una agregación GROUP BY sobre los registros de {parents}.",
        "wd_group_dims":    "Dimensiones de agrupación: {dims}.",
        "wd_group_aggs":    "Agregaciones calculadas: {aggs}.",
        "wd_group_accum":   "La agregación es acumulativa entre períodos.",
        "wd_group_out":     "Alimenta a: {children}.",
        "wd_union_intro":   "{name} fusiona {n} segmento(s) en un dataset unificado.",
        "wd_union_trigger": "Lado trigger: '{seg}' de {res}.",
        "wd_union_extra":   "Fuentes adicionales: {names}.",
        "wd_union_out":     "Alimenta a: {children}.",
        "wd_recon_intro":   "{name} realiza una conciliación estándar entre dos lados:",
        "wd_recon_side":    "Lado {prefix}{trigger}: recurso '{res}', grupo '{seg}'.",
        "wd_recon_rs_n":    "Rule sets de matching: {n}.",
        "wd_recon_rs":      "[{pos}] {name}: {rules}.",
        "wd_recon_chained": "Esta conciliación es encadenada — depende de un paso de matching previo.",
        "wd_recon_out":     "Alimenta a: {children}.",
        "wd_adv_intro":     "{name} realiza una conciliación avanzada con {n} grupo(s):",
        "wd_adv_group":     "Lado {prefix}: recurso '{res}', grupo '{seg}', segmentado por '{col}' ({vals}).",
        "wd_adv_no_int":    "sin segmentación interna",
        "wd_adv_rs_n":      "Rule sets: {n}.",
        "wd_adv_rs":        "[{pos}] {name}: {a} vs {b}.",
        "wd_adv_out":       "Alimenta a: {children}.",
        "wd_generic_intro": "{name} es un recurso de tipo {rt}.",
        "wd_generic_io":    "Entradas: {parents}. Salidas: {children}.",
        "wd_no_parents":    "sin fuentes upstream",
        "wd_no_children":   "sin consumidores downstream",
    },
}

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
C = {
    "red":    "EA0050",
    "red2":   "C0003A",
    "white":  "FFFFFF",
    "grey":   "FDF0F3",
    "grey2":  "F9FAFB",
    "dark":   "1A1A2E",
    "border": "E5E7EB",
    "slate":  "6B7280",
    "blue":   "1D4ED8",
}

RT_COLOR = {
    "native":                  "1D4ED8",
    "source_union":            "0F766E",
    "source_group":            "92400E",
    "reconciliation":          "EA0050",
    "advanced_reconciliation": "6D28D9",
    "consolidation":           "374151",
    "resource_join":           "065F46",
    "cumulative_balance":      "065F46",
}

RT_ORDER = {
    "native": 1, "source_union": 2, "source_group": 3,
    "reconciliation": 4, "advanced_reconciliation": 5,
    "consolidation": 6, "resource_join": 7, "cumulative_balance": 8,
}

def get_rt_label(rt, lang):
    return T[lang].get(f"rt_{rt}", rt)

# ══════════════════════════════════════════════════════════════════════════════
# THEME
# ══════════════════════════════════════════════════════════════════════════════
def get_theme(dark):
    if dark:
        return {
            "bg":          "#0D1117",
            "bg2":         "#161B22",
            "bg3":         "#21262D",
            "text":        "#E6EDF3",
            "text2":       "#8B949E",
            "text3":       "#6E7681",
            "border":      "#30363D",
            "border2":     "#444C56",
            "card_bg":     "#161B22",
            "card_border": "#30363D",
            "input_bg":    "#0D1117",
            "tag_bg":      "#21262D",
            "tag_border":  "#444C56",
        }
    return {
        "bg":          "#F4F5F7",
        "bg2":         "#FFFFFF",
        "bg3":         "#F9FAFB",
        "text":        "#1A1A1A",
        "text2":       "#6B7280",
        "text3":       "#9CA3AF",
        "border":      "#E5E7EB",
        "border2":     "#D1D5DB",
        "card_bg":     "#FFFFFF",
        "card_border": "#E5E7EB",
        "input_bg":    "#FFFFFF",
        "tag_bg":      "#F3F4F6",
        "tag_border":  "#E5E7EB",
    }

def inject_css(TH):
    st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"], .stApp, .main {{
    background-color: {TH['bg']} !important;
    color: {TH['text']} !important;
}}
[data-testid="stHeader"] {{ background-color: transparent !important; }}
.block-container {{ background-color: {TH['bg']} !important; padding-top: 1.2rem !important; }}

html, body, [class*="css"], .stMarkdown, p, span, label, h1, h2, h3, h4,
div[data-testid], .stCaption {{
    font-family: 'Inter', sans-serif !important;
    color: {TH['text']} !important;
}}
code {{ font-family: 'JetBrains Mono', monospace !important; color: #EA0050 !important;
       background: {'#2D1B1B' if TH['bg']=='#0D1117' else '#FEF2F2'} !important; }}

[data-testid="stCheckbox"] label {{ color: {TH['text']} !important; }}
[data-testid="stCheckbox"] > div {{ align-items: center; justify-content: center; height: 100%; margin-top: 25px; }}

div[data-testid="stFileUploader"] section {{
    border: 2px dashed {TH['border2']} !important;
    border-radius: 12px !important;
    background: {TH['bg2']} !important;
}}
div[data-testid="stFileUploader"] section:hover {{ border-color: #EA0050 !important; }}
div[data-testid="stFileUploader"] section button {{
    background: {TH['bg2']} !important;
    color: {TH['text']} !important;
    border: 1px solid {TH['border2']} !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
}}
div[data-testid="stFileUploader"] section button span {{ display: none !important; }}
div[data-testid="stFileUploader"] section button::after {{
    content: 'Browse files';
    font-family: 'Inter', sans-serif;
    font-size: 0.875rem;
    font-weight: 600;
    color: {TH['text']};
}}
div[data-testid="stFileUploader"] small,
div[data-testid="stFileUploader"] p {{ color: {TH['text2']} !important; }}

[data-testid="stMultiSelect"] > div > div {{
    background: {TH['input_bg']} !important;
    border: 1px solid {TH['border2']} !important;
    color: {TH['text']} !important;
    border-radius: 8px !important;
}}
span[data-baseweb="tag"] {{
    background: {TH['tag_bg']} !important;
    color: {TH['text']} !important;
    border-radius: 6px !important;
    border: 1px solid {TH['tag_border']} !important;
}}

.stProgress > div > div {{ background: #EA0050 !important; }}
.stProgress > div {{ background: {TH['border']} !important; }}

div[data-testid="stButton"] button[kind="primary"] {{
    background: #EA0050 !important;
    border: 1px solid #C0003A !important;
    font-size: 1rem !important;
    font-weight: 600 !important;
    color: #ffffff !important;
    border-radius: 8px !important;
    box-shadow: 0 4px 12px rgba(234,0,80,0.2) !important;
}}
div[data-testid="stButton"] button[kind="primary"]:hover {{
    background: #C0003A !important;
}}
div[data-testid="stDownloadButton"] button {{
    background: #EA0050 !important;
    color: #ffffff !important;
    border: 1px solid #C0003A !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
}}

::-webkit-scrollbar {{ width: 8px; height: 8px; }}
::-webkit-scrollbar-track {{ background: {TH['bg']}; }}
::-webkit-scrollbar-thumb {{ background: {TH['border2']}; border-radius: 4px; }}
</style>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# OPENPYXL HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def mk_border():
    t = Side(border_style="thin", color=C["border"])
    return Border(left=t, right=t, top=t, bottom=t)

def sc(cell, bg=None, bold=False, color=C["dark"], size=10,
       ha="left", va="top", wrap=True):
    cell.border = mk_border()
    cell.alignment = Alignment(horizontal=ha, vertical=va, wrap_text=wrap)
    cell.font = Font(name="Calibri", bold=bold, size=size, color=color)
    if bg:
        cell.fill = PatternFill(start_color=bg, end_color=bg, fill_type="solid")

def hdr(cell, text, bg=C["dark"]):
    cell.value = text
    sc(cell, bg=bg, bold=True, color=C["white"], size=10,
       ha="center", va="center", wrap=False)

def section_title(ws, row, text, bg=C["red"], cols=5):
    effective_bg = C["red2"] if bg == C["red"] and text.startswith("  ") else bg
    for i in range(1, cols + 1):
        c = ws.cell(row=row, column=i)
        if i == 1:
            c.value = text
        sc(c, bg=effective_bg, bold=True, color=C["white"], size=10,
           ha="left", va="center", wrap=False)
    ws.merge_cells(f"A{row}:{chr(64+cols)}{row}")
    ws.row_dimensions[row].height = 20
    return row + 1

def meta_row(ws, row, label, value, cols=5, bg_val=None, bg_label=C["slate"]):
    bg_val = bg_val or C["grey"]
    c_l = ws.cell(row=row, column=1)
    c_l.value = label
    sc(c_l, bg=bg_label, bold=True, color=C["white"], size=9, ha="left", va="center", wrap=False)
    val_str = str(value) if value is not None else "—"
    for i in range(2, cols + 1):
        c = ws.cell(row=row, column=i)
        if i == 2:
            c.value = val_str
        sc(c, bg=bg_val, size=9, va="center", wrap=True)
    ws.merge_cells(f"B{row}:{chr(64+cols)}{row}")
    ws.row_dimensions[row].height = 14
    return row + 1

def row_height(n_lines, base=13):
    return max(14, n_lines * base)

# ══════════════════════════════════════════════════════════════════════════════
# PARSERS
# ══════════════════════════════════════════════════════════════════════════════
def build_maps(data):
    res_map, col_map, seg_map, meta_map, seg_usage = {}, {}, {}, {}, {}

    for r in data.get("resources", []):
        eid = r.get("export_id")
        res_map[eid] = r.get("name", str(eid))
        for c in (r.get("columns") or []):
            cid = c.get("export_id")
            col_map[cid] = c.get("label") or c.get("name") or str(cid)
        sg = r.get("source_group") or {}
        for c in sg.get("columns", []):
            cid = c.get("column_id")
            if cid and cid not in col_map:
                col_map[cid] = f"col_{cid}"
        for v in sg.get("values", []):
            cid = v.get("column_id")
            if cid and cid not in col_map:
                col_map[cid] = f"col_{cid}"
        adv = r.get("advanced_reconciliation") or {}
        for rg in adv.get("reconcilable_groups", []):
            for cs in rg.get("columns_selection", []):
                cid = cs.get("column_id")
                if cid and cid not in col_map:
                    col_map[cid] = f"col_{cid}"
            sc2 = rg.get("segmentation_config") or {}
            for m in sc2.get("segmentation_metadata", []):
                meta_map[m.get("export_id")] = m.get("value", "?")
            ccid = sc2.get("criteria_column_id")
            if ccid and ccid not in col_map:
                col_map[ccid] = f"col_{ccid}"
        for seg in (r.get("segments") or []):
            rules = []
            for fset in (seg.get("segment_filter_sets") or []):
                for rule in (fset.get("segment_filter_rules") or []):
                    rules.append(rule)
            seg_map[seg.get("export_id")] = {
                "name": seg.get("name", ""),
                "resource": r.get("name", ""),
                "resource_id": eid,
                "default": seg.get("default_segment", False),
                "rules": rules,
            }

    for r in data.get("resources", []):
        rname = r.get("name", "")
        recon = r.get("reconciliation") or {}
        sa = recon.get("segment_a_id")
        sb = recon.get("segment_b_id")
        pa = recon.get("segment_a_prefix", "A")
        pb = recon.get("segment_b_prefix", "B")
        if sa:
            seg_usage.setdefault(sa, []).append((rname, f"Recon side {pa}"))
        if sb:
            seg_usage.setdefault(sb, []).append((rname, f"Recon side {pb}"))
        adv = r.get("advanced_reconciliation") or {}
        for rg in adv.get("reconcilable_groups", []):
            sgid = rg.get("segment_id")
            if sgid:
                seg_usage.setdefault(sgid, []).append(
                    (rname, f"Adv. Recon. side {rg.get('prefix_side','?')}")
                )
        su = r.get("source_union") or {}
        for us in su.get("union_segments", []):
            sgid = us.get("segment_id")
            if sgid:
                seg_usage.setdefault(sgid, []).append((rname, "Source Union"))

    return res_map, col_map, seg_map, meta_map, seg_usage


def fmt_filter_rules(rules, col_map, no_filters_msg="No filters configured"):
    lines = []
    for r in rules:
        col_name = col_map.get(r.get("column_id"), f"ID:{r.get('column_id')}")
        lines.append(
            f"{r.get('condition','')} [{col_name}] {r.get('operator','')} {r.get('value','')}".strip()
        )
    return "\n".join(lines) if lines else no_filters_msg


def parse_transformation_logic(col, res_map, col_map, S):
    lines = []
    uniq = col.get("uniqueness")
    if uniq:
        dtype = col.get("data_format", "")
        order_keys = uniq.get("order_keys", [])
        part_keys = uniq.get("partition_keys", [])
        if dtype == "boolean":
            lines.append(S["xl_dup_bool"])
        elif dtype == "integer":
            lines.append(S["xl_dup_int"])
        if order_keys:
            order_parts = []
            for ok in sorted(order_keys, key=lambda x: x.get("position", 0)):
                col_name = col_map.get(ok.get("column_id"), f"ID:{ok.get('column_id')}")
                direction = "ASC" if ok.get("order_by", 1) == 1 else "DESC"
                order_parts.append(f"{col_name} {direction}")
            lines.append(S["xl_order_by"] + ", ".join(order_parts))
        if part_keys:
            part_names = [col_map.get(pk.get("column_id"), f"ID:{pk.get('column_id')}")
                          for pk in part_keys]
            lines.append(S["xl_part_by"] + "\n  ".join(part_names))
        return "\n".join(lines)

    v = col.get("v_lookup")
    if v:
        vs = v.get("v_lookup_set") or {}
        origin_id = vs.get("origin_source_id")
        origin = res_map.get(origin_id, f"ID:{origin_id}")
        rules = vs.get("rules", [])
        keys = " & ".join(
            "A." + col_map.get(r.get("column_a_id"), "?") +
            " = B." + col_map.get(r.get("column_b_id"), "?")
            for r in rules
        )
        lines.append(S["xl_vlookup"] + origin)
        if keys:
            lines.append(S["xl_match_key"] + keys)

    parents = [t for t in (col.get("transformations") or []) if t.get("is_parent")]
    for t in parents:
        q = (t.get("query") or "").strip()
        if q and q.upper() != "N/A":
            lines.append(S["xl_formula"] + q)

    return "\n".join(lines) if lines else S["xl_direct"]


def parse_std_reconciliation(recon, res_map, col_map, seg_map, no_filters_msg="No filters configured"):
    if not recon:
        return None
    sa_id = recon.get("segment_a_id")
    sb_id = recon.get("segment_b_id")
    a_cfg = recon.get("a_source_settings") or {}
    b_cfg = recon.get("b_source_settings") or {}

    def resolve_side(cfg, seg_id, prefix):
        resource_name = res_map.get(cfg.get("resource_id"), "—")
        seg = seg_map.get(seg_id) or {}
        seg_name = seg.get("name", f"ID:{seg_id}")
        seg_rules = fmt_filter_rules(seg.get("rules", []), col_map, no_filters_msg)
        return {
            "prefix": prefix, "resource_name": resource_name,
            "group_name": seg_name, "group_filters": seg_rules,
            "is_trigger": cfg.get("is_trigger", False),
        }

    sides = [
        resolve_side(a_cfg, sa_id, recon.get("segment_a_prefix", "A")),
        resolve_side(b_cfg, sb_id, recon.get("segment_b_prefix", "B")),
    ]
    rule_sets = []
    for rs in sorted(recon.get("reconciliation_rule_sets", []),
                     key=lambda x: x.get("position", 99)):
        rules_desc = []
        for rule in (rs.get("reconciliation_rules") or []):
            col_a = col_map.get(rule.get("column_a_id"), f"ID:{rule.get('column_a_id')}")
            col_b = col_map.get(rule.get("column_b_id"), f"ID:{rule.get('column_b_id')}")
            op = rule.get("operator", "=")
            tol = rule.get("tolerance", 0)
            tol_u = rule.get("tolerance_unit") or ""
            tol_s = f"  [tolerance ±{tol} {tol_u}]" if tol else ""
            rules_desc.append(f"A.{col_a}  {op}  B.{col_b}{tol_s}")
        rule_sets.append({
            "pos": rs.get("position", 0), "name": rs.get("name", ""),
            "rules": rules_desc,
            "zero_tol": all(
                rule.get("tolerance", 0) == 0
                for rule in (rs.get("reconciliation_rules") or [])
            ),
        })
    return {"sides": sides, "is_chained": recon.get("is_chained", False), "rule_sets": rule_sets}


def parse_adv_reconciliation(adv, res_map, col_map, seg_map, meta_map, no_filters_msg="No filters configured"):
    if not adv:
        return None
    groups = []
    for rg in (adv.get("reconcilable_groups") or []):
        prefix = rg.get("prefix_side", "?")
        seg_id = rg.get("segment_id")
        seg = seg_map.get(seg_id) or {}
        resource_name = seg.get("resource", res_map.get(rg.get("resource_id"), "—"))
        seg_rules = fmt_filter_rules(seg.get("rules", []), col_map, no_filters_msg)
        sc2 = rg.get("segmentation_config") or {}
        crit_id = sc2.get("criteria_column_id")
        crit_col = col_map.get(crit_id, f"ID:{crit_id}") if crit_id else "—"
        segments = [m.get("value", "") for m in sc2.get("segmentation_metadata", []) if m.get("value")]
        groups.append({
            "prefix": prefix, "resource_name": resource_name,
            "group_name": seg.get("name", f"ID:{seg_id}"),
            "group_filters": seg_rules, "crit_col": crit_col, "segments": segments,
        })
    rule_sets = []
    for rs in sorted(adv.get("reconciliation_rule_sets", []),
                     key=lambda x: x.get("position", 99)):
        rules_desc = []
        for rule in (rs.get("reconciliation_rules") or []):
            col_a = col_map.get(rule.get("column_a_id"), f"ID:{rule.get('column_a_id')}")
            col_b = col_map.get(rule.get("column_b_id"), f"ID:{rule.get('column_b_id')}")
            op = rule.get("operator", "=")
            tol = rule.get("tolerance", 0)
            tol_u = rule.get("tolerance_unit") or ""
            tol_s = f"  [tolerance ±{tol} {tol_u}]" if tol else ""
            rules_desc.append(f"A.{col_a}  {op}  B.{col_b}{tol_s}")
        sweep = []
        for sw in (rs.get("sweep_sides") or []):
            p = sw.get("prefix_side", "?")
            isr = sw.get("input_sweep_resource") or {}
            meta_id = isr.get("segmentation_metadata_id")
            seg_val = meta_map.get(meta_id, f"ID:{meta_id}") if meta_id else "(full resource)"
            sweep.append(f"Side {p}: {seg_val}")
        rule_sets.append({
            "pos": rs.get("position", 0), "name": rs.get("name", ""),
            "cross_type": rs.get("cross_type", ""),
            "new_ver": rs.get("is_new_version", False),
            "rules": rules_desc, "sweep": sweep,
        })
    return {"groups": groups, "rule_sets": rule_sets}


def parse_segment_filters(segs, col_map, no_filters_msg="No filters configured"):
    result = []
    for seg in (segs or []):
        rules = []
        for fset in (seg.get("segment_filter_sets") or []):
            for r in (fset.get("segment_filter_rules") or []):
                col_name = col_map.get(r.get("column_id"), f"ID:{r.get('column_id')}")
                rules.append(
                    f"{r.get('condition','')} [{col_name}] {r.get('operator','')} {r.get('value','')}".strip()
                )
        if rules:
            result.append({"seg_id": seg.get("export_id"), "name": seg.get("name", ""), "rules": rules})
    return result


def parse_source_group(sg, col_map):
    if not sg:
        return [], []
    group_cols = [col_map.get(c.get("column_id"), f"ID:{c.get('column_id')}")
                  for c in sorted(sg.get("columns", []), key=lambda x: x.get("position", 0))]
    agg_vals = [(v.get("function", "?"), col_map.get(v.get("column_id"), f"ID:{v.get('column_id')}"))
                for v in sorted(sg.get("values", []), key=lambda x: x.get("position", 0))]
    return group_cols, agg_vals


def limpiar_hoja(nombre, eid):
    clean = re.sub(r"[\\/*?:\[\]]", "", str(nombre))
    return (clean[:18] + "_" + str(eid))[:31]


def sort_key(r):
    return (RT_ORDER.get(r.get("resource_type", ""), 99), r.get("export_id", 0))


def build_relations(resources, nodes, res_map):
    all_ids = {r.get("export_id") for r in resources}
    rels = {r.get("export_id"): {"parents": [], "children": []} for r in resources}
    for n in nodes:
        t_id = n.get("target")
        s_val = n.get("source")
        if not (t_id and s_val):
            continue
        s_list = s_val if isinstance(s_val, list) else [s_val]
        for sid in s_list:
            ext_a = "" if sid in all_ids else " ↗"
            ext_b = "" if t_id in all_ids else " ↗"
            if t_id in rels:
                rels[t_id]["parents"].append(res_map.get(sid, str(sid)) + ext_a)
            if sid in rels:
                rels[sid]["children"].append(res_map.get(t_id, str(t_id)) + ext_b)
    return rels

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY BUILDER — one-page, management-level
# ══════════════════════════════════════════════════════════════════════════════
def resource_role_sentence(r, col_map, seg_map, meta_map, lang):
    """One short sentence describing what a resource does. No technical detail."""
    rt  = r.get("resource_type", "")
    no_fil = T[lang]["xl_no_filters"]

    if rt == "native":
        return (
            "Ingests raw records directly from the source system."
            if lang == "en" else
            "Ingesta registros directamente desde el sistema fuente."
        )

    if rt == "source_group":
        sg = r.get("source_group") or {}
        group_cols, agg_vals = parse_source_group(sg, col_map)
        dims = ", ".join(group_cols[:3]) + ("…" if len(group_cols) > 3 else "")
        n_agg = len(agg_vals)
        if lang == "en":
            return f"Groups records by {dims or '—'} and computes {n_agg} aggregation(s)."
        return f"Agrupa registros por {dims or '—'} y calcula {n_agg} agregación(es)."

    if rt == "source_union":
        n = len((r.get("source_union") or {}).get("union_segments") or [])
        return (
            f"Merges {n} source segment(s) into a single unified dataset."
            if lang == "en" else
            f"Fusiona {n} segmento(s) en un único dataset unificado."
        )

    if rt == "reconciliation":
        recon = r.get("reconciliation") or {}
        a_cfg = recon.get("a_source_settings") or {}
        b_cfg = recon.get("b_source_settings") or {}
        sa = res_map_placeholder.get(a_cfg.get("resource_id"), "Side A")
        sb = res_map_placeholder.get(b_cfg.get("resource_id"), "Side B")
        n_rs = len(recon.get("reconciliation_rule_sets") or [])
        if lang == "en":
            return f"Matches '{sa}' against '{sb}' using {n_rs} rule set(s)."
        return f"Concilia '{sa}' contra '{sb}' con {n_rs} rule set(s)."

    if rt == "advanced_reconciliation":
        adv = r.get("advanced_reconciliation") or {}
        n_g = len(adv.get("reconcilable_groups") or [])
        n_rs = len(adv.get("reconciliation_rule_sets") or [])
        if lang == "en":
            return f"Advanced cross-match across {n_g} group(s) using {n_rs} rule set(s)."
        return f"Conciliación avanzada de {n_g} grupo(s) con {n_rs} rule set(s)."

    if rt == "source_union":
        n = len((r.get("source_union") or {}).get("union_segments") or [])
        return (
            f"Merges {n} source segment(s) into a unified dataset."
            if lang == "en" else
            f"Fusiona {n} segmento(s) fuente en un dataset unificado."
        )

    if rt == "consolidation":
        return ("Consolidates data from upstream resources."
                if lang == "en" else "Consolida datos de los recursos upstream.")

    if rt == "cumulative_balance":
        return ("Computes a cumulative balance from upstream data."
                if lang == "en" else "Calcula un balance acumulado desde datos upstream.")

    return ("Processes data from upstream resources."
            if lang == "en" else "Procesa datos de recursos upstream.")


# module-level placeholder so resource_role_sentence can access res_map
res_map_placeholder: dict = {}


def build_overview_sentences(resources, rels, res_map, lang):
    """Two-sentence executive overview. No enumeration of technical steps."""
    counts  = Counter(r.get("resource_type", "") for r in resources)
    native  = [r.get("name", "") for r in resources if r.get("resource_type") == "native"]
    recons  = [r.get("name", "") for r in resources
               if r.get("resource_type") in ("reconciliation", "advanced_reconciliation")]
    terminal = [r.get("name", "") for r in resources if not rels[r.get("export_id")]["children"]]
    n_middle = len(resources) - len(native) - len(recons)

    if lang == "en":
        # Sentence 1: what the flow does end-to-end
        src_part  = f"{len(native)} source(s) ({', '.join(native)})" if native else "upstream sources"
        mid_part  = f", transforms data through {n_middle} intermediate step(s)," if n_middle > 0 else ""
        rec_part  = (f" and reconciles it via {', '.join(recons)}" if recons else "")
        s1 = f"This flow ingests data from {src_part}{mid_part}{rec_part}."
        # Sentence 2: output / result
        s2 = (f"The final output is produced by {', '.join(terminal)}."
              if terminal else
              "All resources feed into downstream reconciliation processes.")
    else:
        src_part  = f"{len(native)} fuente(s) ({', '.join(native)})" if native else "fuentes upstream"
        mid_part  = f", lo transforma en {n_middle} paso(s) intermedio(s)," if n_middle > 0 else ""
        rec_part  = (f" y lo concilia a través de {', '.join(recons)}" if recons else "")
        s1 = f"Este flujo ingesta datos desde {src_part}{mid_part}{rec_part}."
        s2 = (f"El output final es producido por {', '.join(terminal)}."
              if terminal else
              "Todos los recursos alimentan procesos de conciliación downstream.")

    return s1, s2


def collect_top_observations(resources, rels, seg_map, seg_usage, lang, max_obs=4):
    """Return at most max_obs high-priority observations. Business language, no tech detail."""
    S   = T[lang]
    obs = []
    seen: set = set()

    def add(priority, text):
        if text not in seen:
            seen.add(text)
            obs.append((priority, text))

    for r in resources:
        eid  = r.get("export_id")
        name = r.get("name", "")
        recon = r.get("reconciliation") or {}

        # Priority 1 — chained reconciliation (operational dependency)
        if recon.get("is_chained"):
            add(1, S["wd_obs_chained"].format(name=name))

        # Priority 2 — zero tolerance (business risk: any mismatch = break)
        for rs in recon.get("reconciliation_rule_sets", []):
            if all(rule.get("tolerance", 0) == 0
                   for rule in (rs.get("reconciliation_rules") or [])):
                add(2, S["wd_obs_zero_tol"].format(rs=rs.get("name", "?"), res=name))
                break   # one alert per resource is enough

        # Priority 3 — unused segments (configuration hygiene)
        for seg in (r.get("segments") or []):
            seg_id = seg.get("export_id")
            if seg_id and not seg_usage.get(seg_id) and not seg.get("default_segment"):
                add(3, S["wd_obs_unused_seg"].format(seg=seg.get("name", "?"), res=name))
                break   # one per resource

    obs.sort(key=lambda x: x[0])
    return [text for _, text in obs[:max_obs]]


# ══════════════════════════════════════════════════════════════════════════════
# WORD GENERATOR — one-page executive summary
# ══════════════════════════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_font(cell, size_pt=9, bold=False,
                   rgb=(30, 30, 30), align=WD_ALIGN_PARAGRAPH.LEFT):
    for para in cell.paragraphs:
        para.alignment = align
        for run in para.runs:
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(*rgb)


def _tight(para, space_before=0, space_after=3):
    """Remove excess paragraph spacing."""
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)


def _section_label(doc, text, rgb=(234, 0, 80)):
    """Small all-caps section label — replaces H1 to save vertical space."""
    p = doc.add_paragraph()
    _tight(p, space_before=10, space_after=3)
    run = p.add_run(text.upper())
    run.font.size  = Pt(8)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*rgb)
    return p


def generar_word(data, selected_ids, lang):
    global res_map_placeholder
    S = T[lang]

    all_resources = data.get("resources", [])
    nodes         = data.get("nodes", [])
    res_map, col_map, seg_map, meta_map, seg_usage = build_maps(data)
    res_map_placeholder = res_map          # expose to resource_role_sentence

    seen, resources = set(), []
    for r in all_resources:
        eid = r.get("export_id")
        if eid in selected_ids and eid not in seen:
            seen.add(eid)
            resources.append(r)
    resources.sort(key=sort_key)
    rels = build_relations(resources, nodes, res_map)

    counts  = Counter(r.get("resource_type", "") for r in resources)
    n_total = len(resources)

    doc = Document()

    # Tight margins to maximise usable space
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── HEADER BAR (colored table: title | date | resource count) ──────────
    hdr_tbl = doc.add_table(rows=1, cols=3)
    hdr_tbl.style = "Table Grid"

    hdr_data = [
        S["wd_title"],
        datetime.now().strftime("%Y-%m-%d  %H:%M"),
        f"{n_total} {S['wd_res_count'].lower()}",
    ]
    hdr_widths = [9, 4, 3]   # approximate column proportions

    for i, (txt, w) in enumerate(zip(hdr_data, hdr_widths)):
        cell = hdr_tbl.rows[0].cells[i]
        cell.text = txt
        _set_cell_bg(cell, "EA0050")
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        _set_cell_font(cell, size_pt=10 if i == 0 else 9,
                       bold=(i == 0), rgb=(255, 255, 255), align=align)

    # ── SUBTITLE ───────────────────────────────────────────────────────────
    p_sub = doc.add_paragraph(S["wd_subtitle"])
    _tight(p_sub, space_before=4, space_after=8)
    for run in p_sub.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run.font.italic = True

    # ── OVERVIEW ───────────────────────────────────────────────────────────
    _section_label(doc, S["wd_s1"])
    s1, s2 = build_overview_sentences(resources, rels, res_map, lang)

    for sentence in (s1, s2):
        p = doc.add_paragraph(sentence)
        _tight(p, space_after=3)
        for run in p.runs:
            run.font.size = Pt(10)

    # ── KEY FIGURES (inline compact metrics) ───────────────────────────────
    p_fig = doc.add_paragraph()
    _tight(p_fig, space_before=6, space_after=6)

    def _kv(label, val):
        """Add 'LABEL value  ' to a paragraph."""
        r1 = p_fig.add_run(f"{label.upper()}: ")
        r1.font.size  = Pt(8.5)
        r1.font.bold  = True
        r1.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        r2 = p_fig.add_run(f"{val}    ")
        r2.font.size  = Pt(9)
        r2.font.bold  = True
        r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    _kv(S["m_total"],   n_total)
    _kv(S["m_sources"], counts.get("native", 0))
    recon_n = counts.get("reconciliation", 0) + counts.get("advanced_reconciliation", 0)
    _kv(S["m_std"] if lang == "en" else "Conc.", recon_n)
    _kv(S["m_groups"],  counts.get("source_group", 0))

    # ── FLOW TABLE: Resource | Type | What it does ─────────────────────────
    _section_label(doc, S["wd_s2"])

    col_headers = S["wd_s2_col"]      # 3-element list now
    tbl = doc.add_table(rows=1 + len(resources), cols=3)
    tbl.style = "Table Grid"

    for i, h_text in enumerate(col_headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h_text
        _set_cell_bg(cell, "1A1A2E")
        _set_cell_font(cell, size_pt=8, bold=True, rgb=(255, 255, 255),
                       align=WD_ALIGN_PARAGRAPH.LEFT)

    for idx, r in enumerate(resources):
        eid  = r.get("export_id")
        rt   = r.get("resource_type", "")
        row  = tbl.rows[idx + 1]
        fill = "F4F5F7" if idx % 2 == 0 else "FFFFFF"
        color_hex = RT_COLOR.get(rt, "374151")
        rt_rgb    = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))

        row.cells[0].text = r.get("name", "")
        row.cells[1].text = get_rt_label(rt, lang)
        row.cells[2].text = resource_role_sentence(r, col_map, seg_map, meta_map, lang)

        for j, cell in enumerate(row.cells):
            _set_cell_bg(cell, fill)
            txt_rgb = rt_rgb if j == 1 else (30, 30, 30)
            bold    = (j == 1)
            _set_cell_font(cell, size_pt=8.5, bold=bold, rgb=txt_rgb)

    # ── OBSERVATIONS (max 4, high priority only) ───────────────────────────
    obs = collect_top_observations(resources, rels, seg_map, seg_usage, lang)
    if obs:
        _section_label(doc, S["wd_s4"])
        for ob in obs:
            p_b  = doc.add_paragraph(style="List Bullet")
            _tight(p_b, space_after=2)
            r_b  = p_b.add_run(ob)
            r_b.font.size = Pt(9)

    # ── FOOTER DISCLAIMER ─────────────────────────────────────────────────
    p_disc = doc.add_paragraph()
    _tight(p_disc, space_before=14, space_after=0)
    run_disc = p_disc.add_run(S["disc_footer_2"])
    run_disc.font.size  = Pt(7.5)
    run_disc.font.italic = True
    run_disc.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ══════════════════════════════════════════════════════════════════════════════
# EXCEL GENERATOR — technical report
# ══════════════════════════════════════════════════════════════════════════════
def generar_excel(data, selected_ids, lang):
    S = T[lang]
    all_resources = data.get("resources", [])
    nodes = data.get("nodes", [])
    res_map, col_map, seg_map, meta_map, seg_usage = build_maps(data)

    seen, resources = set(), []
    for r in all_resources:
        eid = r.get("export_id")
        if eid in selected_ids and eid not in seen:
            seen.add(eid)
            resources.append(r)
    resources.sort(key=sort_key)

    rels = build_relations(resources, nodes, res_map)
    map_hojas = {r.get("export_id"): limpiar_hoja(r.get("name", ""), r.get("export_id"))
                 for r in resources}
    no_fil = S["xl_no_filters"]

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        wb = writer.book

        # ── INDEX ────────────────────────────────────────────────────────────
        ws = wb.create_sheet("📚 Index", 0)
        ws.sheet_view.showGridLines = False

        for i in range(1, 8):
            c = ws.cell(row=1, column=i)
            if i == 1:
                c.value = S["xl_idx_title"]
            sc(c, bg=C["red"], bold=True, color=C["white"], size=13,
               ha="center", va="center", wrap=False)
        ws.merge_cells("A1:G1")
        ws.row_dimensions[1].height = 32

        for i in range(1, 8):
            c = ws.cell(row=2, column=i)
            if i == 1:
                c.value = (f"{S['xl_idx_subtitle']}   |   "
                           f"{S['wd_generated']}: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
                           f"{S['wd_res_count']}: {len(resources)}")
            sc(c, bg=C["dark"], color=C["white"], size=9, ha="center", va="center", wrap=False)
        ws.merge_cells("A2:G2")
        ws.row_dimensions[2].height = 15

        idx_hdrs = [S["xl_col_num"], S["xl_col_id"], S["xl_col_name"], S["xl_col_type"],
                    S["xl_col_from"], S["xl_col_to"], S["xl_col_link"]]
        for i, h in enumerate(idx_hdrs, 1):
            hdr(ws.cell(row=4, column=i), h, bg=C["dark"])
        ws.row_dimensions[4].height = 20
        ws.freeze_panes = "A5"

        for row_n, res in enumerate(resources, 5):
            eid = res.get("export_id")
            rt = res.get("resource_type", "")
            bg = C["grey"] if row_n % 2 == 0 else C["white"]
            vals = [row_n - 4, eid, res.get("name", ""), get_rt_label(rt, lang),
                    ", ".join(rels[eid]["parents"]) or S["xl_origin"],
                    ", ".join(rels[eid]["children"]) or S["xl_end"]]
            for col_n, val in enumerate(vals, 1):
                c = ws.cell(row_n, col_n, val)
                sc(c, bg=bg, size=9, va="center", wrap=False)
                if col_n == 4:
                    c.font = Font(name="Calibri", bold=True, size=9, color=RT_COLOR.get(rt, C["dark"]))
                c.border = mk_border()
            lnk = ws.cell(row_n, 7, S["xl_link_txt"])
            lnk.hyperlink = f"#'{map_hojas[eid]}'!A1"
            lnk.font = Font(name="Calibri", color=C["blue"], underline="single", size=9)
            lnk.border = mk_border()
            ws.row_dimensions[row_n].height = 15

        for col_n, w in enumerate([6, 11, 46, 28, 38, 38, 8], 1):
            ws.column_dimensions[chr(64 + col_n)].width = w

        # ── DETAIL SHEETS ─────────────────────────────────────────────────────
        for res in resources:
            eid = res.get("export_id")
            rt = res.get("resource_type", "")
            name = res.get("name", "")
            tc = RT_COLOR.get(rt, C["dark"])
            COLS = 5

            ws = wb.create_sheet(map_hojas[eid])
            ws.sheet_view.showGridLines = False

            row = 1
            for i in range(1, 6):
                c = ws.cell(row=row, column=i)
                if i == 1:
                    c.value = get_rt_label(rt, lang) + "  ·  " + name
                sc(c, bg=tc, bold=True, color=C["white"], size=12,
                   ha="left", va="center", wrap=False)
            ws.merge_cells(f"A{row}:E{row}")
            ws.row_dimensions[row].height = 30
            row += 1
            ws.freeze_panes = "A2"

            row = meta_row(ws, row, S["xl_meta_id"], eid, cols=COLS, bg_label=tc)
            row = meta_row(ws, row, S["xl_meta_type"], get_rt_label(rt, lang), cols=COLS, bg_label=tc)
            row = meta_row(ws, row, S["xl_meta_from"],
                           ", ".join(rels[eid]["parents"]) or S["xl_origin"], cols=COLS, bg_label=tc)
            row = meta_row(ws, row, S["xl_meta_to"],
                           ", ".join(rels[eid]["children"]) or S["xl_end"], cols=COLS, bg_label=tc)
            row += 1

            # Std reconciliation
            std = parse_std_reconciliation(res.get("reconciliation"), res_map, col_map, seg_map, no_fil)
            if std:
                row = section_title(ws, row, S["xl_s_std"], bg=tc, cols=COLS)
                row = section_title(ws, row, S["xl_s_std_g"], bg=tc, cols=COLS)
                for col_n, h in enumerate([S["xl_h_side"], S["xl_h_resource"],
                                           S["xl_h_group"], S["xl_h_filters"]], 1):
                    hdr(ws.cell(row=row, column=col_n), h, bg=tc)
                hdr(ws.cell(row=row, column=5), "", bg=tc)
                ws.merge_cells(f"D{row}:E{row}")
                ws.row_dimensions[row].height = 18
                row += 1

                for i, side in enumerate(std["sides"]):
                    bg = C["grey"] if i % 2 == 0 else "FFFFFF"
                    trig = S["xl_trigger"] if side["is_trigger"] else ""
                    c1 = ws.cell(row=row, column=1); c1.value = side["prefix"] + trig
                    c2 = ws.cell(row=row, column=2); c2.value = side["resource_name"]
                    c3 = ws.cell(row=row, column=3); c3.value = side["group_name"]
                    c4 = ws.cell(row=row, column=4); c4.value = side["group_filters"]
                    c5 = ws.cell(row=row, column=5); c5.value = ""
                    for c, al in [(c1,"center"),(c2,"left"),(c3,"left"),(c4,"left"),(c5,"left")]:
                        sc(c, bg=bg, size=9, va="top", wrap=True, ha=al)
                    ws.merge_cells(f"D{row}:E{row}")
                    ws.row_dimensions[row].height = row_height(side["group_filters"].count("\n") + 1)
                    row += 1
                row += 1

                row = meta_row(ws, row, S["xl_chained"],
                               S["xl_yes"] if std["is_chained"] else S["xl_no"],
                               cols=COLS, bg_label=tc)
                row += 1

                row = section_title(ws, row, S["xl_s_std_rs"], bg=tc, cols=COLS)
                for col_n, h in enumerate([S["xl_h_pos"], S["xl_h_rsname"], S["xl_h_rules"]], 1):
                    hdr(ws.cell(row=row, column=col_n), h, bg=tc)
                hdr(ws.cell(row=row, column=4), "", bg=tc)
                hdr(ws.cell(row=row, column=5), "", bg=tc)
                ws.merge_cells(f"C{row}:E{row}")
                ws.row_dimensions[row].height = 18
                row += 1

                for i, rs in enumerate(std["rule_sets"]):
                    bg = C["grey"] if i % 2 == 0 else "FFFFFF"
                    c1 = ws.cell(row=row, column=1); c1.value = rs["pos"]
                    c2 = ws.cell(row=row, column=2); c2.value = rs["name"]
                    c3 = ws.cell(row=row, column=3); c3.value = "\n".join(rs["rules"])
                    c4 = ws.cell(row=row, column=4); c4.value = ""
                    c5 = ws.cell(row=row, column=5); c5.value = ""
                    for c, al in [(c1,"center"),(c2,"left"),(c3,"left"),(c4,"left"),(c5,"left")]:
                        sc(c, bg=bg, size=9, va="top", wrap=True, ha=al)
                    ws.merge_cells(f"C{row}:E{row}")
                    ws.row_dimensions[row].height = row_height(len(rs["rules"]))
                    row += 1
                row += 1

            # Adv reconciliation
            adv_parsed = parse_adv_reconciliation(
                res.get("advanced_reconciliation"), res_map, col_map, seg_map, meta_map, no_fil)
            if adv_parsed:
                row = section_title(ws, row, S["xl_s_adv"], bg=tc, cols=COLS)
                row = section_title(ws, row, S["xl_s_adv_g"], bg=tc, cols=COLS)
                for col_n, h in enumerate([S["xl_h_side"], S["xl_h_resource"],
                                           S["xl_h_group"], S["xl_h_filters"], S["xl_h_intseg"]], 1):
                    hdr(ws.cell(row=row, column=col_n), h, bg=tc)
                ws.row_dimensions[row].height = 18
                row += 1

                for i, g in enumerate(adv_parsed["groups"]):
                    bg = C["grey"] if i % 2 == 0 else "FFFFFF"
                    segs_txt = "\n".join(g["segments"]) if g["segments"] else S["xl_no_intseg"]
                    n_lines = max(g["group_filters"].count("\n") + 1, len(g["segments"]) or 1)
                    c1 = ws.cell(row=row, column=1); c1.value = g["prefix"]
                    c2 = ws.cell(row=row, column=2); c2.value = g["resource_name"]
                    c3 = ws.cell(row=row, column=3); c3.value = g["group_name"]
                    c4 = ws.cell(row=row, column=4); c4.value = g["group_filters"]
                    c5 = ws.cell(row=row, column=5); c5.value = segs_txt
                    for c, al in [(c1,"center"),(c2,"left"),(c3,"left"),(c4,"left"),(c5,"left")]:
                        sc(c, bg=bg, size=9, va="top", wrap=True, ha=al)
                    ws.row_dimensions[row].height = row_height(n_lines)
                    row += 1
                row += 1

                row = section_title(ws, row, S["xl_s_adv_rs"], bg=tc, cols=COLS)
                for col_n, h in enumerate([S["xl_h_pos"], S["xl_h_rsname"], S["xl_h_rules"],
                                           S["xl_h_seg_a"], S["xl_h_seg_b"]], 1):
                    hdr(ws.cell(row=row, column=col_n), h, bg=tc)
                ws.row_dimensions[row].height = 18
                row += 1

                for i, rs in enumerate(adv_parsed["rule_sets"]):
                    bg = C["grey"] if i % 2 == 0 else "FFFFFF"
                    name_txt = rs["name"]
                    if rs["cross_type"]:
                        name_txt += "\n[" + rs["cross_type"] + "]"
                    if rs["new_ver"]:
                        name_txt += "  ✦ new version"
                    seg_a = next((s.replace("Side A: ", "") for s in rs["sweep"] if "A:" in s), "—")
                    seg_b = next((s.replace("Side B: ", "") for s in rs["sweep"] if "B:" in s), "—")
                    c1 = ws.cell(row=row, column=1); c1.value = rs["pos"]
                    c2 = ws.cell(row=row, column=2); c2.value = name_txt
                    c3 = ws.cell(row=row, column=3); c3.value = "\n".join(rs["rules"])
                    c4 = ws.cell(row=row, column=4); c4.value = seg_a
                    c5 = ws.cell(row=row, column=5); c5.value = seg_b
                    for c, al in [(c1,"center"),(c2,"left"),(c3,"left"),(c4,"left"),(c5,"left")]:
                        sc(c, bg=bg, size=9, va="top", wrap=True, ha=al)
                    ws.row_dimensions[row].height = row_height(max(len(rs["rules"]), 1))
                    row += 1
                row += 1

            # Group By
            sg = res.get("source_group")
            if sg:
                row = section_title(ws, row, S["xl_s_gb"], bg=tc, cols=COLS)
                group_cols, agg_vals = parse_source_group(sg, col_map)
                row = meta_row(ws, row, S["xl_h_gbdims"],
                               " | ".join(group_cols) or "—", cols=COLS, bg_label=tc)
                agg_str = "  |  ".join(f"{fn}( {col} )" for fn, col in agg_vals)
                row = meta_row(ws, row, S["xl_h_aggs"], agg_str or "—", cols=COLS, bg_label=tc)
                row = meta_row(ws, row, S["xl_h_accum"],
                               S["xl_yes"] if sg.get("is_accumulative") else S["xl_no"],
                               cols=COLS, bg_label=tc)
                row += 1

            # Source Union
            su = res.get("source_union")
            if su:
                row = section_title(ws, row, S["xl_s_union"], bg=tc, cols=COLS)
                for col_n, h in enumerate([S["xl_h_source"], S["xl_h_group"],
                                           S["xl_h_role"], S["xl_h_filters"]], 1):
                    hdr(ws.cell(row=row, column=col_n), h, bg=tc)
                hdr(ws.cell(row=row, column=5), "", bg=tc)
                ws.merge_cells(f"D{row}:E{row}")
                ws.row_dimensions[row].height = 18
                row += 1

                for i, us in enumerate(su.get("union_segments") or []):
                    seg_id = us.get("segment_id")
                    seg_info = seg_map.get(seg_id) or {}
                    resource_name = seg_info.get("resource", f"ID:{seg_id}")
                    group_name = seg_info.get("name", f"ID:{seg_id}")
                    filters_text = fmt_filter_rules(seg_info.get("rules", []), col_map, no_fil)
                    rol = S["xl_trig_src"] + (us.get("trigger_type") or "") if us.get("is_trigger") else S["xl_add_src"]
                    bg = C["grey"] if i % 2 == 0 else "FFFFFF"
                    c1 = ws.cell(row=row, column=1); c1.value = resource_name
                    c2 = ws.cell(row=row, column=2); c2.value = group_name
                    c3 = ws.cell(row=row, column=3); c3.value = rol
                    c4 = ws.cell(row=row, column=4); c4.value = filters_text
                    c5 = ws.cell(row=row, column=5); c5.value = ""
                    for c, al in [(c1,"left"),(c2,"left"),(c3,"center"),(c4,"left"),(c5,"left")]:
                        sc(c, bg=bg, size=9, va="top", wrap=True, ha=al)
                    ws.merge_cells(f"D{row}:E{row}")
                    ws.row_dimensions[row].height = row_height(max(filters_text.count("\n") + 1, 1))
                    row += 1
                row += 1

            # Segments
            segs_all = parse_segment_filters(res.get("segments", []), col_map, no_fil)
            if segs_all:
                row = section_title(ws, row, S["xl_s_segs"], bg=tc, cols=COLS)
                for col_n, h in enumerate([S["xl_h_gname"], S["xl_h_fapplied"]], 1):
                    hdr(ws.cell(row=row, column=col_n), h, bg=tc)
                hdr(ws.cell(row=row, column=3), "", bg=tc)
                hdr(ws.cell(row=row, column=4), "", bg=tc)
                hdr(ws.cell(row=row, column=5), S["xl_h_usedin"], bg=tc)
                ws.merge_cells(f"B{row}:D{row}")
                ws.row_dimensions[row].height = 18
                row += 1
                for i, seg in enumerate(segs_all):
                    bg = C["grey"] if i % 2 == 0 else "FFFFFF"
                    usages = seg_usage.get(seg["seg_id"], [])
                    usage_text = ("\n".join(u[0] + " (" + u[1] + ")" for u in usages)
                                  if usages else S["xl_no_usage"])
                    n_lines = max(len(seg["rules"]), len(usages) if usages else 1)
                    c1 = ws.cell(row=row, column=1); c1.value = seg["name"]
                    c2 = ws.cell(row=row, column=2); c2.value = "\n".join(seg["rules"])
                    c3 = ws.cell(row=row, column=3); c3.value = ""
                    c4 = ws.cell(row=row, column=4); c4.value = ""
                    c5 = ws.cell(row=row, column=5); c5.value = usage_text
                    for c in [c1, c2, c3, c4]:
                        sc(c, bg=bg, size=9, va="top", wrap=True)
                    sc(c5, bg=bg, size=9, va="top", wrap=True,
                       color="365C42" if usages else "4B5563")
                    ws.merge_cells(f"B{row}:D{row}")
                    ws.row_dimensions[row].height = row_height(n_lines)
                    row += 1
                row += 1

            # Columns
            columns = sorted(res.get("columns") or [], key=lambda x: x.get("position", 0))
            if columns:
                row = section_title(ws, row, S["xl_s_cols"], bg=tc, cols=COLS)
                for col_n, h in enumerate([S["xl_h_label"], S["xl_h_dtype"],
                                           S["xl_h_ctype"], S["xl_h_logic"]], 1):
                    hdr(ws.cell(row=row, column=col_n), h, bg=tc)
                hdr(ws.cell(row=row, column=5), "", bg=tc)
                ws.merge_cells(f"D{row}:E{row}")
                ws.row_dimensions[row].height = 18
                row += 1
                for i, col in enumerate(columns):
                    label = col.get("label") or col.get("name", "")
                    dtype = col.get("data_format", "")
                    col_type = (col.get("column_type") or "").replace("_", " ").upper()
                    logic = parse_transformation_logic(col, res_map, col_map, S)
                    bg = C["grey"] if i % 2 == 0 else "FFFFFF"
                    c1 = ws.cell(row=row, column=1); c1.value = label
                    c2 = ws.cell(row=row, column=2); c2.value = dtype
                    c3 = ws.cell(row=row, column=3); c3.value = col_type
                    c4 = ws.cell(row=row, column=4); c4.value = logic
                    c5 = ws.cell(row=row, column=5); c5.value = ""
                    for c, al in [(c1,"left"),(c2,"center"),(c3,"center"),(c4,"left"),(c5,"left")]:
                        sc(c, bg=bg, size=9, va="top", wrap=True, ha=al)
                    ws.merge_cells(f"D{row}:E{row}")
                    ws.row_dimensions[row].height = row_height(logic.count("\n") + 1)
                    row += 1

            ws.column_dimensions["A"].width = 26
            ws.column_dimensions["B"].width = 22
            ws.column_dimensions["C"].width = 22
            ws.column_dimensions["D"].width = 22
            ws.column_dimensions["E"].width = 36

        if "Sheet" in wb.sheetnames:
            wb.remove(wb["Sheet"])

    output.seek(0)
    return output

# ══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ══════════════════════════════════════════════════════════════════════════════
lang = st.session_state.lang
dark = st.session_state.dark
S = T[lang]
TH = get_theme(dark)
inject_css(TH)

# ── TOP BAR: language + theme ─────────────────────────────────────────────────
top_l, top_m, top_r = st.columns([2, 2, 6])

with top_l:
    new_lang = st.radio(
        S["lang_label"],
        options=["en", "es"],
        format_func=lambda x: "🇬🇧 English" if x == "en" else "🇦🇷 Español",
        index=0 if lang == "en" else 1,
        horizontal=True,
        key="lang_radio",
    )
    if new_lang != st.session_state.lang:
        st.session_state.lang = new_lang
        st.rerun()

with top_m:
    dark_toggle = st.toggle(S["theme_label"], value=dark, key="dark_toggle")
    if dark_toggle != st.session_state.dark:
        st.session_state.dark = dark_toggle
        st.rerun()

# ── HEADER ────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div style='background:linear-gradient(135deg,#EA0050 0%,#C0003A 100%);padding:28px 32px;
    border-radius:12px;margin-bottom:32px;box-shadow:0 8px 20px -5px rgba(234,0,80,0.3)'>
  <div style='color:#FFF;font-family:Inter,sans-serif;font-size:1.75rem;font-weight:700;letter-spacing:-0.5px'>
    {S['header_title']}
  </div>
  <div style='color:rgba(255,255,255,0.85);font-family:Inter,sans-serif;font-size:0.95rem;margin-top:6px;font-weight:500'>
    {S['header_sub']}
  </div>
</div>""", unsafe_allow_html=True)

# ── UPLOAD ────────────────────────────────────────────────────────────────────
up = st.file_uploader(
    S["upload_label"],
    type=["json"],
    help=S["upload_help"],
    label_visibility="visible",
)

if not up:
    st.markdown(f"""
<div style='background:{TH['card_bg']};border:1px solid {TH['card_border']};border-radius:12px;
    padding:48px 32px;text-align:center;margin-top:16px'>
  <div style='font-size:2.5rem;margin-bottom:12px'>📂</div>
  <p style='color:{TH['text']};font-size:1.1rem;font-weight:600;margin:0'>{S['upload_placeholder_h']}</p>
  <p style='color:{TH['text2']};font-size:0.9rem;margin:8px 0 0'>{S['upload_placeholder_p']}</p>
</div>""", unsafe_allow_html=True)
    # Privacy notice on empty state
    st.markdown(
        f"<div style='background:{'#0D2B0D' if dark else '#F0FDF4'};border:1px solid {'#1A4A1A' if dark else '#BBF7D0'};"
        f"border-radius:10px;padding:14px 18px;margin-top:16px;display:flex;gap:12px;align-items:flex-start'>"
        f"<div><p style='margin:0;font-size:0.85rem;font-weight:600;color:{'#4ADE80' if dark else '#166534'}'>"
        f"{S['disc_privacy_title']}</p>"
        f"<p style='margin:4px 0 0;font-size:0.8rem;color:{'#86EFAC' if dark else '#15803D'};line-height:1.5'>"
        f"{S['disc_privacy_body']}</p></div></div>",
        unsafe_allow_html=True,
    )
    st.markdown(
        f"<p style='text-align:center;color:{TH['text2']};font-size:0.78rem;margin-top:24px'>{S['footer']}</p>",
        unsafe_allow_html=True,
    )
    st.stop()

try:
    data = json.load(up)
except Exception as e:
    st.error(f"Error reading JSON: {e}")
    st.stop()

all_resources = data.get("resources", [])
nodes = data.get("nodes", [])

seen_load, resources_unique = set(), []
for r in all_resources:
    eid = r.get("export_id")
    if eid not in seen_load:
        seen_load.add(eid)
        resources_unique.append(r)

res_map, col_map, seg_map, meta_map, seg_usage = build_maps(data)
resources_unique.sort(key=sort_key)
rels_all = build_relations(resources_unique, nodes, res_map)

# ── METRICS ───────────────────────────────────────────────────────────────────
_tc = Counter(r.get("resource_type", "") for r in resources_unique)
_total = len(resources_unique)
_nombre = up.name if len(up.name) <= 30 else up.name[:27] + "…"

def _card(label, value, color="#EA0050"):
    return (
        f"<div style='background:{TH['card_bg']};border:1px solid {TH['card_border']};"
        f"border-radius:12px;box-shadow:0 2px 4px rgba(0,0,0,0.04);padding:16px;text-align:center'>"
        f"<div style='font-size:0.72rem;color:{TH['text2']};font-family:Inter,sans-serif;"
        f"font-weight:600;margin-bottom:6px;text-transform:uppercase;letter-spacing:0.5px'>{label}</div>"
        f"<div style='font-size:1.7rem;font-weight:700;color:{color};font-family:Inter,sans-serif;"
        f"line-height:1'>{value}</div></div>"
    )

_cards = (
    f"<div style='display:grid;grid-template-columns:repeat(7,1fr);gap:12px;margin-bottom:8px'>"
    + _card(S["m_total"], _total, TH["text"])
    + _card(S["m_sources"], _tc.get("native", 0), f"#{RT_COLOR['native']}")
    + _card(S["m_unions"], _tc.get("source_union", 0), f"#{RT_COLOR['source_union']}")
    + _card(S["m_groups"], _tc.get("source_group", 0), f"#{RT_COLOR['source_group']}")
    + _card(S["m_std"], _tc.get("reconciliation", 0), f"#{RT_COLOR['reconciliation']}")
    + _card(S["m_adv"], _tc.get("advanced_reconciliation", 0), f"#{RT_COLOR['advanced_reconciliation']}")
    + (f"<div style='background:{TH['card_bg']};border:1px solid {TH['card_border']};"
       f"border-radius:12px;padding:16px;text-align:left;display:flex;flex-direction:column;justify-content:center'>"
       f"<div style='font-size:0.72rem;color:{TH['text2']};font-weight:600;text-transform:uppercase;"
       f"letter-spacing:0.5px;margin-bottom:6px'>{S['m_json']}</div>"
       f"<div style='font-size:0.82rem;font-weight:600;color:{TH['text']};word-break:break-all;line-height:1.4'>"
       f"{_nombre}</div></div>")
    + "</div>"
)
st.markdown(_cards, unsafe_allow_html=True)
st.markdown(f"<hr style='margin:28px 0;border-color:{TH['border']}'>", unsafe_allow_html=True)

# ── STEP 1: SELECTION ─────────────────────────────────────────────────────────
st.markdown(f"<h3 style='margin-bottom:16px;font-weight:700;color:{TH['text']}'>{S['step1_title']}</h3>",
            unsafe_allow_html=True)

all_types = sorted({r.get("resource_type", "") for r in resources_unique},
                   key=lambda x: RT_ORDER.get(x, 99))

col_f1, col_f2 = st.columns([4, 1])
with col_f1:
    filtro_tipo = st.multiselect(
        "filter",
        options=all_types,
        format_func=lambda x: get_rt_label(x, lang),
        default=all_types,
        label_visibility="collapsed",
        placeholder=S["filter_placeholder"],
    )

resources_visible = [r for r in resources_unique if r.get("resource_type", "") in filtro_tipo]

bc1, bc2, bc3 = st.columns([1, 1, 6])
select_all = bc1.button(S["btn_all"], use_container_width=True)
deselect_all = bc2.button(S["btn_none"], use_container_width=True)

if "sel" not in st.session_state:
    st.session_state.sel = {r.get("export_id"): True for r in resources_unique}
if select_all:
    for r in resources_visible:
        st.session_state.sel[r.get("export_id")] = True
if deselect_all:
    for r in resources_visible:
        st.session_state.sel[r.get("export_id")] = False

st.write("")

tipo_groups: dict = {}
for r in resources_visible:
    tipo_groups.setdefault(r.get("resource_type", ""), []).append(r)

selected_ids = set()
for rt in sorted(tipo_groups.keys(), key=lambda x: RT_ORDER.get(x, 99)):
    group = tipo_groups[rt]
    color_hex = RT_COLOR.get(rt, "475569")
    label = get_rt_label(rt, lang)

    st.markdown(
        f"<div style='display:flex;align-items:center;gap:12px;margin:36px 0 16px'>"
        f"<span style='background:#{color_hex}18;color:#{color_hex};border:1px solid #{color_hex}40;"
        f"padding:6px 14px;border-radius:20px;font-size:0.85rem;font-weight:700;white-space:nowrap'>"
        f"{label}</span>"
        f"<span style='color:{TH['text2']};font-size:0.85rem;font-weight:600'>"
        f"{len(group)} {S['resources_n']}</span></div>",
        unsafe_allow_html=True,
    )

    for r in group:
        eid = r.get("export_id")
        name = r.get("name", "")
        pars = ", ".join(rels_all[eid]["parents"]) or S["origin"]
        chils = ", ".join(rels_all[eid]["children"]) or S["end_flow"]

        ca, cb = st.columns([0.3, 9.7])
        checked = ca.checkbox("", value=st.session_state.sel.get(eid, True), key=f"chk_{eid}")
        st.session_state.sel[eid] = checked

        opacity = "1" if checked else "0.5"
        bg_color = TH["card_bg"] if checked else TH["bg3"]
        border_col = "#EA0050" if checked else TH["card_border"]
        border_w = "2px" if checked else "1px"

        cb.markdown(
            f"<div style='opacity:{opacity};padding:14px 18px;background:{bg_color};border-radius:8px;"
            f"margin-bottom:8px;border:{border_w} solid {border_col};transition:all 0.15s'>"
            f"<div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:6px'>"
            f"<span style='font-weight:700;font-size:0.95rem;color:{TH['text']}'>{name}</span>"
            f"<span style='font-size:0.72rem;color:{TH['text2']};font-family:JetBrains Mono,monospace;"
            f"background:{TH['bg3']};padding:4px 8px;border-radius:4px;border:1px solid {TH['border']}'>{eid}</span>"
            f"</div>"
            f"<div style='font-size:0.8rem;color:{TH['text2']};display:flex;gap:24px;font-weight:500'>"
            f"<span>← <span style='color:{TH['text3']}'>{pars[:75]}{'…' if len(pars)>75 else ''}</span></span>"
            f"<span>→ <span style='color:{TH['text3']}'>{chils[:75]}{'…' if len(chils)>75 else ''}</span></span>"
            f"</div></div>",
            unsafe_allow_html=True,
        )
        if checked:
            selected_ids.add(eid)

# ── STEP 2: GENERATE ──────────────────────────────────────────────────────────
st.markdown(f"<hr style='margin:48px 0 24px;border-color:{TH['border']}'>", unsafe_allow_html=True)

n_sel = len(selected_ids)
if n_sel == 0:
    st.warning(S["no_sel_warning"])
    st.stop()

# Selection summary badges
tipos_sel = Counter(r.get("resource_type", "") for r in resources_unique if r.get("export_id") in selected_ids)
badges = "".join(
    f"<span style='background:#{RT_COLOR.get(rt,'6B7280')}15;color:#{RT_COLOR.get(rt,'6B7280')};"
    f"border:1px solid #{RT_COLOR.get(rt,'6B7280')}30;padding:4px 12px;border-radius:12px;"
    f"font-size:0.8rem;font-weight:700;white-space:nowrap'>{get_rt_label(rt,lang)} ({cnt})</span> "
    for rt, cnt in sorted(tipos_sel.items(), key=lambda x: RT_ORDER.get(x[0], 99))
)
st.markdown(
    f"<div style='background:{TH['card_bg']};border:1px solid {TH['card_border']};border-radius:12px;"
    f"padding:16px 20px;margin-bottom:20px;display:flex;align-items:center;gap:12px;flex-wrap:wrap'>"
    f"<span style='font-weight:700;color:{TH['text']};white-space:nowrap;font-size:1.05rem'>"
    f"📋 {n_sel} {S['sel_label']}:</span>{badges}</div>",
    unsafe_allow_html=True,
)

base_name = os.path.splitext(up.name)[0]
ts = datetime.now().strftime("%Y-%m-%d_%H%M")
fname_excel = f"flowdocs_{base_name}_{ts}.xlsx"
fname_word = f"flowdocs_{base_name}_{ts}.docx"

# Output disclaimer — shown before the generate button
_disc_bg  = "#1A1500" if dark else "#FFFBEB"
_disc_bdr = "#3D3000" if dark else "#FDE68A"
_disc_h   = "#FCD34D" if dark else "#92400E"
_disc_p   = "#FDE68A" if dark else "#78350F"
st.markdown(
    f"<div style='background:{_disc_bg};border:1px solid {_disc_bdr};"
    f"border-radius:10px;padding:14px 18px;margin-bottom:16px;display:flex;gap:12px;align-items:flex-start'>"
    f"<div><p style='margin:0;font-size:0.85rem;font-weight:600;color:{_disc_h}'>"
    f"{S['disc_output_title']}</p>"
    f"<p style='margin:4px 0 0;font-size:0.8rem;color:{_disc_p};line-height:1.5'>"
    f"{S['disc_output_body']}</p></div></div>",
    unsafe_allow_html=True,
)

if st.button(S["btn_generate"], type="primary", use_container_width=True):
    prog = st.progress(0, text="Starting…")
    try:
        prog.progress(20, text="Parsing resources…")
        excel_bytes = generar_excel(data, selected_ids, lang)
        prog.progress(60, text="Building executive summary…")
        word_bytes = generar_word(data, selected_ids, lang)
        prog.progress(100, text="Done.")

        st.success(S["success_msg"].format(n=n_sel))

        dl1, dl2 = st.columns(2)
        dl1.download_button(
            label=S["btn_excel"],
            data=excel_bytes,
            file_name=fname_excel,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            type="primary",
        )
        dl2.download_button(
            label=S["btn_word"],
            data=word_bytes,
            file_name=fname_word,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )
    except Exception as e:
        prog.empty()
        st.error(f"Error generating reports: {e}")
        import traceback
        st.code(traceback.format_exc())

st.markdown(
    f"<hr style='margin:32px 0;border-color:{TH['border']}'>"
    f"<div style='text-align:center;padding-bottom:24px'>"
    f"<p style='color:{TH['text2']};font-size:0.78rem;margin:0 0 4px'>Simetrik Flow Docs · v3.0 &nbsp;·&nbsp; MIT License</p>"
    f"<p style='color:{TH['text3']};font-size:0.75rem;margin:0 0 3px'>{S['disc_footer_1']}</p>"
    f"<p style='color:{TH['text3']};font-size:0.75rem;margin:0 0 3px'>{S['disc_footer_2']}</p>"
    f"<p style='color:{TH['text3']};font-size:0.75rem;margin:0'>{S['disc_footer_3']}</p>"
    f"</div>",
    unsafe_allow_html=True,
)
